import pandas as pd
import numpy as np
from io import BytesIO
import streamlit as st
from datetime import datetime, date
import json
from typing import Dict, List, Any, Optional
import base64

# Excel 관련
import xlsxwriter
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows

# PDF 관련
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Word 문서 관련
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 차트 관련
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots

class RiskAssessmentReportGenerator:
    """위험성평가 보고서 생성기"""
    
    def __init__(self, data_handler=None):
        self.data_handler = data_handler
        self.report_date = datetime.now()
        
    def generate_excel_report(self, data: Dict[str, Any]) -> bytes:
        """엑셀 보고서 생성"""
        buffer = BytesIO()
        
        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
            workbook = writer.book
            
            # 스타일 정의
            header_format = workbook.add_format({
                'bold': True,
                'text_wrap': True,
                'valign': 'center',
                'fg_color': '#D7E4BD',
                'border': 1
            })
            
            cell_format = workbook.add_format({
                'text_wrap': True,
                'valign': 'top',
                'border': 1
            })
            
            # 1. 표지 시트
            self._create_cover_sheet(writer, data, header_format, cell_format)
            
            # 2. 작업공정 시트
            self._create_process_sheet(writer, data, header_format, cell_format)
            
            # 3. 위험정보 시트
            self._create_risk_info_sheet(writer, data, header_format, cell_format)
            
            # 4. 위험성평가 시트
            self._create_assessment_sheet(writer, data, header_format, cell_format)
            
            # 5. 통계 및 차트 시트
            self._create_statistics_sheet(writer, data, header_format, cell_format)
            
        buffer.seek(0)
        return buffer.getvalue()
    
    def _create_cover_sheet(self, writer, data, header_format, cell_format):
        """표지 시트 생성"""
        worksheet = writer.book.add_worksheet('표지')
        
        # 제목
        title_format = writer.book.add_format({
            'bold': True,
            'font_size': 18,
            'align': 'center',
            'valign': 'vcenter'
        })
        
        worksheet.merge_range('A1:D3', f'{self.report_date.year}년도 위험성평가 결과서', title_format)
        
        # 기본 정보
        title_data = data.get('title_data', {})
        info_data = [
            ['회사명', title_data.get('company_name', '')],
            ['주소', title_data.get('company_address', '')],
            ['전화번호', title_data.get('company_tel', '')],
            ['대표자', title_data.get('ceo_name', '')],
            ['작성일', self.report_date.strftime('%Y-%m-%d')]
        ]
        
        row = 5
        for label, value in info_data:
            worksheet.write(row, 0, label, header_format)
            worksheet.write(row, 1, value, cell_format)
            row += 1
            
        # 열 너비 조정
        worksheet.set_column('A:A', 15)
        worksheet.set_column('B:D', 25)
    
    def _create_process_sheet(self, writer, data, header_format, cell_format):
        """작업공정 시트 생성"""
        process_data = data.get('process_data', {})
        processes = process_data.get('processes', [])
        
        if not processes:
            return
            
        df_processes = pd.DataFrame(processes)
        df_processes.to_excel(writer, sheet_name='작업공정', index=False, startrow=1)
        
        worksheet = writer.sheets['작업공정']
        
        # 헤더 스타일 적용
        for col_num, value in enumerate(df_processes.columns.values):
            worksheet.write(1, col_num, value, header_format)
            
        # 열 너비 조정
        worksheet.set_column('A:A', 15)
        worksheet.set_column('B:B', 30)
        worksheet.set_column('C:E', 20)
    
    def _create_risk_info_sheet(self, writer, data, header_format, cell_format):
        """위험정보 시트 생성"""
        risk_data = data.get('risk_info_data', {})
        
        worksheet = writer.book.add_worksheet('위험정보')
        
        # 기본 정보
        worksheet.write(0, 0, '업종명', header_format)
        worksheet.write(0, 1, risk_data.get('industry_name', ''), cell_format)
        
        # 원재료, 장비, 화학물질 정보
        sections = [
            ('원재료', risk_data.get('materials', [])),
            ('장비', risk_data.get('equipment', [])),
            ('화학물질', risk_data.get('chemicals', []))
        ]
        
        row = 2
        for section_name, items in sections:
            worksheet.write(row, 0, section_name, header_format)
            if items:
                for i, item in enumerate(items):
                    worksheet.write(row + i, 1, item, cell_format)
                row += len(items) + 1
            else:
                row += 2
    
    def _create_assessment_sheet(self, writer, data, header_format, cell_format):
        """위험성평가 시트 생성"""
        assessment_data = data.get('assessment_data', [])
        
        if not assessment_data:
            return
            
        df_assessment = pd.DataFrame(assessment_data)
        df_assessment.to_excel(writer, sheet_name='위험성평가표', index=False, startrow=1)
        
        worksheet = writer.sheets['위험성평가표']
        
        # 위험성 레벨에 따른 색상 적용
        high_risk_format = writer.book.add_format({
            'bg_color': '#FFB3BA',
            'border': 1
        })
        medium_risk_format = writer.book.add_format({
            'bg_color': '#FFDFBA',
            'border': 1
        })
        low_risk_format = writer.book.add_format({
            'bg_color': '#BAFFC9',
            'border': 1
        })
        
        # 위험성 점수에 따른 조건부 서식
        for row_num, row_data in enumerate(assessment_data, start=2):
            risk_score = row_data.get('risk_score', 0)
            if isinstance(risk_score, (int, float)):
                if risk_score >= 12:
                    format_to_use = high_risk_format
                elif risk_score >= 6:
                    format_to_use = medium_risk_format
                else:
                    format_to_use = low_risk_format
                    
                # 위험성 점수 셀에 색상 적용
                risk_col = list(df_assessment.columns).index('risk_score') if 'risk_score' in df_assessment.columns else -1
                if risk_col >= 0:
                    worksheet.write(row_num, risk_col, risk_score, format_to_use)
    
    def _create_statistics_sheet(self, writer, data, header_format, cell_format):
        """통계 시트 생성"""
        assessment_data = data.get('assessment_data', [])
        
        if not assessment_data:
            return
            
        worksheet = writer.book.add_worksheet('통계')
        
        # 위험성 통계 계산
        risk_stats = self._calculate_risk_statistics(assessment_data)
        
        # 통계 테이블 작성
        worksheet.write(0, 0, '위험성 평가 통계', header_format)
        
        stats_data = [
            ['구분', '수량', '비율'],
            ['전체 항목', risk_stats['total'], '100%'],
            ['높음 (12-20)', risk_stats['high'], f"{risk_stats['high_percent']:.1f}%"],
            ['보통 (6-11)', risk_stats['medium'], f"{risk_stats['medium_percent']:.1f}%"],
            ['낮음 (1-5)', risk_stats['low'], f"{risk_stats['low_percent']:.1f}%"]
        ]
        
        for row_num, row_data in enumerate(stats_data):
            for col_num, cell_data in enumerate(row_data):
                format_to_use = header_format if row_num == 0 else cell_format
                worksheet.write(row_num + 2, col_num, cell_data, format_to_use)
    
    def _calculate_risk_statistics(self, assessment_data: List[Dict]) -> Dict[str, Any]:
        """위험성 통계 계산"""
        total = len(assessment_data)
        high = sum(1 for item in assessment_data if item.get('risk_score', 0) >= 12)
        medium = sum(1 for item in assessment_data if 6 <= item.get('risk_score', 0) < 12)
        low = sum(1 for item in assessment_data if 1 <= item.get('risk_score', 0) < 6)
        
        return {
            'total': total,
            'high': high,
            'medium': medium,
            'low': low,
            'high_percent': (high / total * 100) if total > 0 else 0,
            'medium_percent': (medium / total * 100) if total > 0 else 0,
            'low_percent': (low / total * 100) if total > 0 else 0
        }
    
    def generate_pdf_report(self, data: Dict[str, Any]) -> bytes:
        """PDF 보고서 생성"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # 스타일 설정
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # 중앙 정렬
        )
        
        story = []
        
        # 제목 페이지
        title = Paragraph(f"{self.report_date.year}년도 위험성평가 결과서", title_style)
        story.append(title)
        story.append(Spacer(1, 12))
        
        # 회사 정보
        title_data = data.get('title_data', {})
        company_info = [
            ['항목', '내용'],
            ['회사명', title_data.get('company_name', '')],
            ['주소', title_data.get('company_address', '')],
            ['전화번호', title_data.get('company_tel', '')],
            ['대표자', title_data.get('ceo_name', '')],
            ['작성일', self.report_date.strftime('%Y-%m-%d')]
        ]
        
        company_table = Table(company_info)
        company_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(company_table)
        story.append(PageBreak())
        
        # 위험성평가 결과
        assessment_data = data.get('assessment_data', [])
        if assessment_data:
            story.append(Paragraph("위험성평가 결과", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            # 통계 차트 (텍스트 형태로)
            risk_stats = self._calculate_risk_statistics(assessment_data)
            stats_text = f"""
            전체 평가 항목: {risk_stats['total']}개
            - 높음 위험: {risk_stats['high']}개 ({risk_stats['high_percent']:.1f}%)
            - 보통 위험: {risk_stats['medium']}개 ({risk_stats['medium_percent']:.1f}%)
            - 낮음 위험: {risk_stats['low']}개 ({risk_stats['low_percent']:.1f}%)
            """
            story.append(Paragraph(stats_text, styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_word_report(self, data: Dict[str, Any]) -> bytes:
        """Word 보고서 생성"""
        doc = Document()
        
        # 제목
        title = doc.add_heading(f'{self.report_date.year}년도 위험성평가 결과서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 회사 정보
        doc.add_heading('1. 사업장 정보', level=1)
        title_data = data.get('title_data', {})
        
        company_table = doc.add_table(rows=5, cols=2)
        company_table.style = 'Table Grid'
        
        company_info = [
            ('회사명', title_data.get('company_name', '')),
            ('주소', title_data.get('company_address', '')),
            ('전화번호', title_data.get('company_tel', '')),
            ('대표자', title_data.get('ceo_name', '')),
            ('작성일', self.report_date.strftime('%Y-%m-%d'))
        ]
        
        for i, (label, value) in enumerate(company_info):
            company_table.cell(i, 0).text = label
            company_table.cell(i, 1).text = value
        
        # 작업공정 정보
        process_data = data.get('process_data', {})
        processes = process_data.get('processes', [])
        
        if processes:
            doc.add_heading('2. 작업공정', level=1)
            for i, process in enumerate(processes, 1):
                doc.add_heading(f'공정 {i}: {process.get("name", "")}', level=2)
                doc.add_paragraph(f'설명: {process.get("description", "")}')
                if process.get('equipment'):
                    doc.add_paragraph(f'장비: {", ".join(process["equipment"])}')
                if process.get('chemicals'):
                    doc.add_paragraph(f'화학물질: {", ".join(process["chemicals"])}')
        
        # 위험성평가 결과
        assessment_data = data.get('assessment_data', [])
        if assessment_data:
            doc.add_heading('3. 위험성평가 결과', level=1)
            
            # 통계
            risk_stats = self._calculate_risk_statistics(assessment_data)
            doc.add_paragraph(f'전체 평가 항목: {risk_stats["total"]}개')
            doc.add_paragraph(f'높음 위험: {risk_stats["high"]}개 ({risk_stats["high_percent"]:.1f}%)')
            doc.add_paragraph(f'보통 위험: {risk_stats["medium"]}개 ({risk_stats["medium_percent"]:.1f}%)')
            doc.add_paragraph(f'낮음 위험: {risk_stats["low"]}개 ({risk_stats["low_percent"]:.1f}%)')
        
        # 바이트로 변환
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def create_risk_chart(self, assessment_data: List[Dict]) -> bytes:
        """위험성 평가 차트 생성"""
        if not assessment_data:
            return None
            
        risk_stats = self._calculate_risk_statistics(assessment_data)
        
        # 파이 차트 생성
        labels = ['높음', '보통', '낮음']
        values = [risk_stats['high'], risk_stats['medium'], risk_stats['low']]
        colors_list = ['#ff6b6b', '#ffd43b', '#51cf66']
        
        fig = go.Figure(data=[go.Pie(
            labels=labels, 
            values=values,
            marker_colors=colors_list,
            textinfo='label+percent',
            textfont_size=12
        )])
        
        fig.update_layout(
            title="위험성 평가 결과 분포",
            font=dict(size=14),
            showlegend=True
        )
        
        # 이미지로 변환
        img_bytes = fig.to_image(format="png", width=800, height=600)
        return img_bytes
    
    def create_process_chart(self, process_data: Dict) -> bytes:
        """공정별 위험요인 차트 생성"""
        processes = process_data.get('processes', [])
        if not processes:
            return None
            
        process_names = [p.get('name', f'공정{i+1}') for i, p in enumerate(processes)]
        equipment_counts = [len(p.get('equipment', [])) for p in processes]
        chemical_counts = [len(p.get('chemicals', [])) for p in processes]
        
        fig = go.Figure()
        
        fig.add_trace(go.Bar(
            name='장비수',
            x=process_names,
            y=equipment_counts,
            marker_color='lightblue'
        ))
        
        fig.add_trace(go.Bar(
            name='화학물질수',
            x=process_names,
            y=chemical_counts,
            marker_color='lightcoral'
        ))
        
        fig.update_layout(
            title='공정별 위험요인 현황',
            xaxis_title='공정명',
            yaxis_title='개수',
            barmode='group',
            font=dict(size=12)
        )
        
        img_bytes = fig.to_image(format="png", width=800, height=500)
        return img_bytes
    
    def get_download_link(self, file_data: bytes, filename: str, file_type: str) -> str:
        """다운로드 링크 생성"""
        b64 = base64.b64encode(file_data).decode()
        
        mime_types = {
            'excel': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'pdf': 'application/pdf',
            'word': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        
        mime_type = mime_types.get(file_type, 'application/octet-stream')
        
        href = f'<a href="data:{mime_type};base64,{b64}" download="{filename}">📥 {filename} 다운로드</a>'
        return href

def create_comprehensive_report(data: Dict[str, Any], report_type: str = 'excel') -> bytes:
    """통합 보고서 생성 함수"""
    generator = RiskAssessmentReportGenerator()
    
    if report_type == 'excel':
        return generator.generate_excel_report(data)
    elif report_type == 'pdf':
        return generator.generate_pdf_report(data)
    elif report_type == 'word':
        return generator.generate_word_report(data)
    else:
        raise ValueError(f"지원하지 않는 보고서 타입: {report_type}")

# Streamlit 사용을 위한 헬퍼 함수들
def display_report_options(data: Dict[str, Any]):
    """Streamlit에서 보고서 옵션 표시"""
    st.subheader("📊 보고서 생성")
    
    col1, col2, col3 = st.columns(3)
    
    generator = RiskAssessmentReportGenerator()
    
    with col1:
        if st.button("📗 엑셀 보고서 생성"):
            try:
                excel_data = generator.generate_excel_report(data)
                filename = f"위험성평가_보고서_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                
                st.download_button(
                    label="📥 엑셀 파일 다운로드",
                    data=excel_data,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
                st.success("엑셀 보고서가 생성되었습니다!")
            except Exception as e:
                st.error(f"엑셀 보고서 생성 중 오류: {e}")
    
    with col2:
        if st.button("📕 PDF 보고서 생성"):
            try:
                pdf_data = generator.generate_pdf_report(data)
                filename = f"위험성평가_보고서_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                
                st.download_button(
                    label="📥 PDF 파일 다운로드",
                    data=pdf_data,
                    file_name=filename,
                    mime="application/pdf"
                )
                st.success("PDF 보고서가 생성되었습니다!")
            except Exception as e:
                st.error(f"PDF 보고서 생성 중 오류: {e}")
    
    with col3:
        if st.button("📘 Word 보고서 생성"):
            try:
                word_data = generator.generate_word_report(data)
                filename = f"위험성평가_보고서_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                
                st.download_button(
                    label="📥 Word 파일 다운로드",
                    data=word_data,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.success("Word 보고서가 생성되었습니다!")
            except Exception as e:
                st.error(f"Word 보고서 생성 중 오류: {e}")

def display_charts(data: Dict[str, Any]):
    """차트 표시"""
    st.subheader("📈 통계 차트")
    
    generator = RiskAssessmentReportGenerator()
    assessment_data = data.get('assessment_data', [])
    
    if assessment_data:
        col1, col2 = st.columns(2)
        
        with col1:
            try:
                risk_chart = generator.create_risk_chart(assessment_data)
                if risk_chart:
                    st.image(risk_chart, caption="위험성 평가 결과 분포")
            except Exception as e:
                st.error(f"위험성 차트 생성 중 오류: {e}")
        
        with col2:
            try:
                process_data = data.get('process_data', {})
                process_chart = generator.create_process_chart(process_data)
                if process_chart:
                    st.image(process_chart, caption="공정별 위험요인 현황")
            except Exception as e:
                st.error(f"공정 차트 생성 중 오류: {e}")
    else:
        st.info("차트를 생성하려면 위험성평가 데이터가 필요합니다.")
